import os
import logging
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import json
from pathlib import Path

import streamlit as st
from groq import Groq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.schema import Document
import PyPDF2
import docx
from datetime import datetime

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('rag_assistant.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

@dataclass
class RAGConfig:
    """Configuration settings for the RAG assistant."""
    chunk_size: int = 1000
    chunk_overlap: int = 200
    max_retrieval_docs: int = 5
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2"
    groq_model: str = "llama-3.3-70b-versatile"
    temperature: float = 0.1
    max_tokens: int = 1000

class DocumentProcessor:
    """Handles document loading and preprocessing."""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""],
            length_function=len,
        )
    
    def load_documents(self, file_paths: List[str]) -> List[Document]:
        """Load documents from various file formats."""
        documents = []
        
        for file_path in file_paths:
            try:
                path = Path(file_path)
                if not path.exists():
                    logger.warning(f"File not found: {file_path}")
                    continue
                
                content = self._extract_text_from_file(path)
                if content:
                    doc = Document(
                        page_content=content,
                        metadata={
                            "source": str(path),
                            "filename": path.name,
                            "file_type": path.suffix,
                            "processed_at": datetime.now().isoformat()
                        }
                    )
                    documents.append(doc)
                    logger.info(f"Loaded document: {path.name}")
                
            except Exception as e:
                logger.error(f"Error loading {file_path}: {str(e)}")
        
        return documents
    
    def _extract_text_from_file(self, file_path: Path) -> Optional[str]:
        """Extract text content from different file types."""
        try:
            if file_path.suffix.lower() == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_path.suffix.lower() in ['.docx', '.doc']:
                return self._extract_from_docx(file_path)
            elif file_path.suffix.lower() == '.txt':
                return self._extract_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_path.suffix}")
                return None
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    def _extract_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT files."""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read().strip()
    
    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """Split documents into smaller chunks."""
        chunked_docs = []
        for doc in documents:
            chunks = self.text_splitter.split_text(doc.page_content)
            for i, chunk in enumerate(chunks):
                chunk_doc = Document(
                    page_content=chunk,
                    metadata={
                        **doc.metadata,
                        "chunk_id": i,
                        "total_chunks": len(chunks)
                    }
                )
                chunked_docs.append(chunk_doc)
        
        logger.info(f"Created {len(chunked_docs)} document chunks")
        return chunked_docs

class VectorStoreManager:
    """Manages the vector store for document retrieval."""
    
    def __init__(self, config: RAGConfig):
        self.config = config
        self.embeddings = HuggingFaceEmbeddings(
            model_name=config.embedding_model,
            model_kwargs={'device': 'cpu'}
        )
        self.vector_store = None
    
    def create_vector_store(self, documents: List[Document]) -> None:
        """Create and populate the vector store."""
        try:
            if not documents:
                raise ValueError("No documents provided for vector store creation")
            
            self.vector_store = FAISS.from_documents(documents, self.embeddings)
            logger.info(f"Created vector store with {len(documents)} documents")
        except Exception as e:
            logger.error(f"Error creating vector store: {str(e)}")
            raise
    
    def save_vector_store(self, path: str) -> None:
        """Save the vector store to disk."""
        if self.vector_store:
            self.vector_store.save_local(path)
            logger.info(f"Vector store saved to {path}")
    
    def load_vector_store(self, path: str) -> None:
        """Load the vector store from disk."""
        try:
            self.vector_store = FAISS.load_local(path, self.embeddings)
            logger.info(f"Vector store loaded from {path}")
        except Exception as e:
            logger.error(f"Error loading vector store: {str(e)}")
            raise
    
    def similarity_search(self, query: str) -> List[Document]:
        """Retrieve relevant documents based on similarity search."""
        if not self.vector_store:
            raise ValueError("Vector store not initialized")
        
        try:
            docs = self.vector_store.similarity_search(
                query, 
                k=self.config.max_retrieval_docs
            )
            logger.info(f"Retrieved {len(docs)} documents for query")
            return docs
        except Exception as e:
            logger.error(f"Error in similarity search: {str(e)}")
            return []

class GroqLLMClient:
    """Client for interacting with Groq API."""
    
    def __init__(self, config: RAGConfig, api_key: str):
        self.config = config
        self.client = Groq(api_key=api_key)
    
    def generate_response(self, query: str, context_docs: List[Document]) -> str:
        """Generate response using retrieved context."""
        try:
            context = self._format_context(context_docs)
            prompt = self._create_prompt(query, context)
            
            response = self.client.chat.completions.create(
                model=self.config.groq_model,
                messages=[{"role": "user", "content": prompt}],
                temperature=self.config.temperature,
                max_tokens=self.config.max_tokens
            )
            
            answer = response.choices[0].message.content
            logger.info("Generated response successfully")
            return answer
            
        except Exception as e:
            logger.error(f"Error generating response: {str(e)}")
            return f"Sorry, I encountered an error generating the response: {str(e)}"
    
    def _format_context(self, docs: List[Document]) -> str:
        """Format retrieved documents into context string."""
        context_parts = []
        for i, doc in enumerate(docs, 1):
            source = doc.metadata.get('filename', 'Unknown')
            content = doc.page_content[:500] + "..." if len(doc.page_content) > 500 else doc.page_content
            context_parts.append(f"Source {i} ({source}):\n{content}")
        
        return "\n\n".join(context_parts)
    
    def _create_prompt(self, query: str, context: str) -> str:
        """Create the prompt for the LLM."""
        return f"""You are a helpful AI assistant that answers questions based on the provided context documents. 

Context:
{context}

Question: {query}

Instructions:
- Answer the question using only the information provided in the context
- If the context doesn't contain enough information to answer the question, say so clearly
- Cite which source(s) you used in your answer
- Be concise and accurate
- If multiple sources provide different information, mention the discrepancy

Answer:"""

class RAGAssistant:
    """Main RAG assistant orchestrator."""
    
    def __init__(self, groq_api_key: str, config: Optional[RAGConfig] = None):
        self.config = config or RAGConfig()
        self.doc_processor = DocumentProcessor(self.config)
        self.vector_manager = VectorStoreManager(self.config)
        self.llm_client = GroqLLMClient(self.config, groq_api_key)
        self.is_initialized = False
    
    def initialize_knowledge_base(self, document_paths: List[str], save_path: Optional[str] = None) -> bool:
        """Initialize the knowledge base from documents."""
        try:
            # Load and process documents
            documents = self.doc_processor.load_documents(document_paths)
            if not documents:
                logger.error("No documents were successfully loaded")
                return False
            
            # Chunk documents
            chunked_docs = self.doc_processor.chunk_documents(documents)
            
            # Create vector store
            self.vector_manager.create_vector_store(chunked_docs)
            
            # Save vector store if path provided
            if save_path:
                self.vector_manager.save_vector_store(save_path)
            
            self.is_initialized = True
            logger.info("Knowledge base initialized successfully")
            return True
            
        except Exception as e:
            logger.error(f"Error initializing knowledge base: {str(e)}")
            return False
    
    def load_knowledge_base(self, vector_store_path: str) -> bool:
        """Load existing knowledge base."""
        try:
            self.vector_manager.load_vector_store(vector_store_path)
            self.is_initialized = True
            logger.info("Knowledge base loaded successfully")
            return True
        except Exception as e:
            logger.error(f"Error loading knowledge base: {str(e)}")
            return False
    
    def ask_question(self, question: str) -> Dict[str, Any]:
        """Process a question and return response with metadata."""
        if not self.is_initialized:
            return {
                "answer": "Knowledge base not initialized. Please load documents first.",
                "sources": [],
                "error": True
            }
        
        try:
            # Retrieve relevant documents
            retrieved_docs = self.vector_manager.similarity_search(question)
            
            if not retrieved_docs:
                return {
                    "answer": "I couldn't find any relevant information to answer your question.",
                    "sources": [],
                    "error": False
                }
            
            # Generate response
            answer = self.llm_client.generate_response(question, retrieved_docs)
            
            # Prepare sources information
            sources = [
                {
                    "filename": doc.metadata.get('filename', 'Unknown'),
                    "chunk_id": doc.metadata.get('chunk_id', 0),
                    "content_preview": doc.page_content[:200] + "..."
                }
                for doc in retrieved_docs
            ]
            
            return {
                "answer": answer,
                "sources": sources,
                "error": False
            }
            
        except Exception as e:
            logger.error(f"Error processing question: {str(e)}")
            return {
                "answer": f"An error occurred: {str(e)}",
                "sources": [],
                "error": True
            }

def main():
    """Main function to run the Streamlit app."""
    st.set_page_config(
        page_title="RAG Assistant",
        page_icon="🤖",
        layout="wide"
    )
    
    st.title("🤖 RAG-Powered Question Answering Assistant")
    st.markdown("Ask questions about your documents using advanced retrieval and AI generation!")
    
    # Sidebar for configuration
    with st.sidebar:
        st.header("Configuration")
        
        # API Key input
        groq_api_key = st.text_input(
            "Groq API Key", 
            type="password",
            help="Enter your Groq API key to use the assistant"
        )
        
        if not groq_api_key:
            st.warning("Please enter your Groq API key to continue")
            st.stop()
        
        # File uploader
        st.subheader("Upload Documents")
        uploaded_files = st.file_uploader(
            "Choose files",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx']
        )
        
        # Advanced settings
        with st.expander("Advanced Settings"):
            chunk_size = st.slider("Chunk Size", 500, 2000, 1000)
            chunk_overlap = st.slider("Chunk Overlap", 50, 500, 200)
            max_docs = st.slider("Max Retrieved Documents", 3, 10, 5)
            temperature = st.slider("Temperature", 0.0, 1.0, 0.1)
    
    # Initialize session state
    if "rag_assistant" not in st.session_state:
        st.session_state.rag_assistant = None
    if "knowledge_base_ready" not in st.session_state:
        st.session_state.knowledge_base_ready = False
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Document Processing")
        
        if uploaded_files and st.button("Process Documents"):
            with st.spinner("Processing documents..."):
                # Save uploaded files temporarily
                temp_paths = []
                for uploaded_file in uploaded_files:
                    temp_path = f"temp_{uploaded_file.name}"
                    with open(temp_path, "wb") as f:
                        f.write(uploaded_file.getvalue())
                    temp_paths.append(temp_path)
                
                # Create RAG assistant with custom config
                config = RAGConfig(
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                    max_retrieval_docs=max_docs,
                    temperature=temperature
                )
                
                rag_assistant = RAGAssistant(groq_api_key, config)
                
                # Initialize knowledge base
                success = rag_assistant.initialize_knowledge_base(temp_paths)
                
                # Clean up temp files
                for temp_path in temp_paths:
                    try:
                        os.remove(temp_path)
                    except:
                        pass
                
                if success:
                    st.session_state.rag_assistant = rag_assistant
                    st.session_state.knowledge_base_ready = True
                    st.success("Documents processed successfully!")
                else:
                    st.error("Failed to process documents. Check the logs for details.")
    
    with col2:
        st.subheader("Status")
        if st.session_state.knowledge_base_ready:
            st.success("✅ Knowledge base ready")
        else:
            st.info("📁 Upload documents to get started")
    
    # Question answering section
    if st.session_state.knowledge_base_ready:
        st.subheader("Ask Questions")
        
        # Sample questions
        with st.expander("Sample Questions"):
            st.markdown("""
            Try asking questions like:
            - What is the main topic of the documents?
            - Summarize the key findings
            - What methodology was used?
            - What are the limitations mentioned?
            """)
        
        question = st.text_input(
            "Your Question:",
            placeholder="Enter your question about the uploaded documents..."
        )
        
        if question and st.button("Get Answer"):
            with st.spinner("Generating answer..."):
                result = st.session_state.rag_assistant.ask_question(question)
                
                if not result["error"]:
                    st.subheader("Answer")
                    st.write(result["answer"])
                    
                    if result["sources"]:
                        st.subheader("Sources")
                        for i, source in enumerate(result["sources"], 1):
                            with st.expander(f"Source {i}: {source['filename']}"):
                                st.write(source["content_preview"])
                else:
                    st.error(result["answer"])

if __name__ == "__main__":
    main()